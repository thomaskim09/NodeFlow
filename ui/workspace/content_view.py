from PySide6.QtWidgets import (
    QWidget,
    QVBoxLayout,
    QHBoxLayout,
    QComboBox,
    QPushButton,
    QTextEdit,
    QFileDialog,
    QMessageBox,
    QDialog,
    QLabel,
    QDialogButtonBox,
    QApplication,
    QFrame,
    QStackedLayout,
    QInputDialog,
)
from PySide6.QtCore import Qt, Signal
from PySide6.QtGui import (
    QTextCursor,
    QColor,
    QTextCharFormat,
    QTextDocument,
    QFont,
)
import os
import database
import docx

from managers.theme_manager import load_settings
from .excel_import_dialog import ExcelImportDialog
from managers import excel_import_manager
from qt_material_icons import MaterialIcon


class ContentView(QWidget):
    document_deleted = Signal()
    document_added = Signal(int)
    bulk_documents_added = Signal()
    segment_clicked = Signal(int)
    text_selection_changed = Signal(bool)
    node_clicked_in_content = Signal(int)
    participant_highlight_requested = Signal(int)
    documents_changed = Signal()
    segments_changed = Signal()

    def __init__(self, project_id):
        super().__init__()
        self.project_id = project_id
        self.documents_map = {}
        self.current_document_id = None
        self.current_participant_id = None
        self.is_dirty = False
        self._coded_segments_cache = []
        self._pending_highlight = None
        self.setAcceptDrops(True)
        main_layout = QVBoxLayout(self)
        top_bar_layout = QHBoxLayout()
        title_label = QLabel("Document View")
        font = title_label.font()
        font.setBold(True)
        title_label.setFont(font)
        self.doc_selector = QComboBox()
        self.doc_selector.setMinimumWidth(300)

        self.import_button = QPushButton()
        import_icon = MaterialIcon("upload")
        self.import_button.setIcon(import_icon)
        self.import_button.setToolTip("Import Document (.txt, .docx, .xlsx)")
        self.save_button = QPushButton()
        save_icon = MaterialIcon("save")
        self.save_button.setIcon(save_icon)
        self.save_button.setToolTip("Save Changes")
        self.save_button.setFixedSize(28, 28)
        self.save_button.setEnabled(False)
        delete_button = QPushButton()
        delete_icon = MaterialIcon("delete")
        delete_button.setIcon(delete_icon)
        delete_button.setToolTip("Delete Current Document")
        delete_button.setFixedSize(28, 28)

        self.text_edit = QTextEdit()
        self.text_edit.setReadOnly(False)
        self.text_edit.setAcceptDrops(False)
        self.drop_overlay = QFrame()
        self.drop_overlay.setObjectName("dropOverlay")
        overlay_layout = QVBoxLayout(self.drop_overlay)
        overlay_layout.setAlignment(Qt.AlignmentFlag.AlignCenter)
        overlay_layout.setSpacing(10)
        icon_label = QLabel()
        upload_icon = MaterialIcon("upload")
        icon_label.setPixmap(upload_icon.pixmap(48, 48))
        icon_label.setAlignment(Qt.AlignmentFlag.AlignCenter)

        text_label = QLabel("Drop document file(s) here\n(.txt, .docx, .xlsx)")
        text_font = QFont()
        text_font.setPointSize(12)
        text_label.setFont(text_font)
        text_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        overlay_layout.addWidget(icon_label)
        overlay_layout.addWidget(text_label)
        self.stacked_layout = QStackedLayout()
        self.stacked_layout.addWidget(self.text_edit)
        self.stacked_layout.addWidget(self.drop_overlay)
        self.stacked_layout.setCurrentWidget(self.text_edit)
        info_bar = QFrame()
        info_bar.setFrameShape(QFrame.Shape.StyledPanel)
        info_bar_layout = QHBoxLayout(info_bar)
        info_bar_layout.setContentsMargins(5, 2, 5, 2)
        self.word_count_label = QLabel("Word Count: 0")
        self.segment_count_label = QLabel("Coded Segments: 0")
        info_bar_layout.addWidget(self.word_count_label)
        info_bar_layout.addStretch()
        info_bar_layout.addWidget(self.segment_count_label)
        top_bar_layout.addWidget(title_label)
        top_bar_layout.addWidget(self.doc_selector)
        top_bar_layout.addStretch()
        top_bar_layout.addWidget(self.import_button)
        top_bar_layout.addWidget(self.save_button)
        top_bar_layout.addWidget(delete_button)
        main_layout.addLayout(top_bar_layout)
        main_layout.addLayout(self.stacked_layout)
        main_layout.addWidget(info_bar)

        self.import_button.clicked.connect(self.open_import_dialog)
        self.save_button.clicked.connect(self.save_document)
        delete_button.clicked.connect(self.delete_current_document)
        self.doc_selector.currentIndexChanged.connect(self.handle_document_switch)
        self.text_edit.textChanged.connect(self.on_text_changed)
        self.text_edit.cursorPositionChanged.connect(self.on_cursor_position_changed)
        self.text_edit.selectionChanged.connect(self.on_selection_changed_for_coding)
        self.load_document_list()

    def _select_and_scroll(self, start, end):
        cursor = self.text_edit.textCursor()
        cursor.setPosition(start)
        cursor.setPosition(end, QTextCursor.MoveMode.KeepAnchor)
        self.text_edit.setTextCursor(cursor)
        self.text_edit.ensureCursorVisible()

    def go_to_segment(self, document_id, start, end):
        if document_id == self.current_document_id:
            self._select_and_scroll(start, end)
        else:
            self._pending_highlight = (start, end)
            id_to_display_text = {v: k for k, v in self.documents_map.items()}
            display_text = id_to_display_text.get(document_id)
            if display_text:
                index = self.doc_selector.findText(display_text)
                if index != -1:
                    self.doc_selector.setCurrentIndex(index)

    def open_import_dialog(self):
        file_paths, _ = QFileDialog.getOpenFileNames(
            self,
            "Import Documents",
            "",
            "All Supported Files (*.txt *.docx *.xlsx);;Text Files (*.txt);;Word Documents (*.docx);;Excel Files (*.xlsx)",
        )
        if file_paths:
            self.handle_files_dropped(file_paths)

    def handle_file_import(self, file_path):
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".xlsx":
            self._import_from_excel(file_path)
        elif ext in [".txt", ".docx"]:
            self._process_text_document(file_path)
        else:
            QMessageBox.warning(
                self, "Unsupported File", f"The file type '{ext}' is not supported."
            )

    def _import_from_excel(self, file_path):
        participants = database.get_participants_for_project(self.project_id)
        dialog = ExcelImportDialog(file_path, participants, self)
        if not dialog.valid_headers:
            QMessageBox.critical(
                self,
                "Error",
                f"Could not read headers from '{os.path.basename(file_path)}'. Please ensure it is a valid .xlsx file with a header row.",
            )
            return
        if dialog.exec() == QDialog.DialogCode.Accepted:
            QApplication.setOverrideCursor(Qt.CursorShape.WaitCursor)
            mappings = dialog.get_column_mappings()
            docs_imported, errors = excel_import_manager.import_data(
                self.project_id, file_path, mappings
            )
            QApplication.restoreOverrideCursor()
            if docs_imported > 0:
                self.bulk_documents_added.emit()
            summary_message = f"Successfully imported {docs_imported} document(s) from '{os.path.basename(file_path)}'."
            if errors:
                detailed_errors = "\n".join(errors[:5])
                if len(errors) > 5:
                    detailed_errors += "\n(And more...)"
                error_dialog = QMessageBox(self)
                error_dialog.setWindowTitle("Import Complete with Errors")
                error_dialog.setText(summary_message)
                error_dialog.setDetailedText(detailed_errors)
                error_dialog.exec()
            else:
                QMessageBox.information(self, "Import Complete", summary_message)

    def dragEnterEvent(self, event):
        mime_data = event.mimeData()
        if mime_data.hasUrls():
            for url in mime_data.urls():
                if url.isLocalFile():
                    file_path = url.toLocalFile().lower()
                    if file_path.endswith((".txt", ".docx", ".xlsx")):
                        event.acceptProposedAction()
                        self.show_drop_overlay()
                        return
        event.ignore()

    def dropEvent(self, event):
        self.hide_drop_overlay()
        if event.mimeData().hasUrls():
            file_paths = [
                url.toLocalFile()
                for url in event.mimeData().urls()
                if url.isLocalFile()
            ]
            self.handle_files_dropped(file_paths)
            event.acceptProposedAction()
        else:
            event.ignore()

    def handle_files_dropped(self, file_paths):
        for path in file_paths:
            self.handle_file_import(path)

    def _process_text_document(self, file_path):
        try:
            content = ""
            if file_path.lower().endswith(".docx"):
                doc = docx.Document(file_path)
                content = "\n\n".join([p.text for p in doc.paragraphs])
            else:
                with open(file_path, "r", encoding="utf-8") as f:
                    content = f.read()
            title = os.path.basename(file_path)
            if database.check_document_exists(self.project_id, title, content):
                reply = QMessageBox.question(
                    self,
                    "Duplicate Document",
                    f"The document '{title}' has been imported before.\n\nDo you want to add it as a new copy?",
                    QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                    QMessageBox.StandardButton.No,
                )
                if reply == QMessageBox.StandardButton.No:
                    return
            participants = database.get_participants_for_project(self.project_id)
            if not participants:
                name, ok = QInputDialog.getText(
                    self,
                    "Create Participant",
                    "Enter a name for the first participant:",
                )
                if not ok or not name.strip():
                    QMessageBox.warning(
                        self,
                        "No Participants",
                        "You must create at least one participant to import documents.",
                    )
                    return
                database.add_participant(self.project_id, name.strip())
                participants = database.get_participants_for_project(self.project_id)
                if not participants:
                    QMessageBox.critical(
                        self,
                        "Error",
                        "Failed to create participant. Please try again.",
                    )
                    return
            if len(participants) == 1:
                self._import_and_add_document(participants[0]["id"], title, content)
            else:
                dialog = AssignParticipantDialog(participants, self)
                if dialog.exec() == QDialog.DialogCode.Accepted:
                    participant_id = dialog.get_selected_participant_id()
                    if participant_id:
                        self._import_and_add_document(participant_id, title, content)
        except Exception as e:
            QMessageBox.critical(
                self,
                "Error",
                f"Failed to process file '{os.path.basename(file_path)}': {e}",
            )

    def dragLeaveEvent(self, event):
        self.hide_drop_overlay()
        event.accept()

    def show_drop_overlay(self):
        if self.stacked_layout.currentWidget() is not self.drop_overlay:
            settings = load_settings()
            is_dark = settings.get("theme") == "Dark"
            bg_color_str = (
                "rgba(60, 60, 60, 0.95)" if is_dark else "rgba(240, 240, 240, 0.95)"
            )
            border_color_str = "#aaa" if is_dark else "#888"
            self.drop_overlay.setStyleSheet(
                f"#dropOverlay {{ background-color: {bg_color_str}; border: 2px dashed {border_color_str}; border-radius: 10px; }}"
            )
            self.stacked_layout.setCurrentWidget(self.drop_overlay)

    def hide_drop_overlay(self):
        self.stacked_layout.setCurrentWidget(self.text_edit)

    def on_selection_changed_for_coding(self):
        self.text_selection_changed.emit(self.text_edit.textCursor().hasSelection())

    def load_document_list(self, doc_id_to_select=None):
        self.doc_selector.blockSignals(True)
        self.doc_selector.clear()
        docs = database.get_documents_for_project(self.project_id)
        display_count = {}
        display_texts = []
        for doc in docs:
            key = (doc["title"], doc["participant_name"] or "Unassigned")
            display_count[key] = display_count.get(key, 0) + 1
        seen = {}
        self.documents_map = {}
        for doc in docs:
            title = doc["title"]
            participant = doc["participant_name"] or "Unassigned"
            key = (title, participant)
            seen[key] = seen.get(key, 0) + 1
            if display_count[key] > 1:
                display_text = f"{title} ({participant}) [{seen[key]}]"
            else:
                display_text = f"{title} ({participant})"
            self.documents_map[display_text] = doc["id"]
            display_texts.append(display_text)
        for display_text in sorted(display_texts):
            self.doc_selector.addItem(display_text)
        self.doc_selector.blockSignals(False)
        new_index = -1
        if doc_id_to_select:
            id_to_display_text = {v: k for k, v in self.documents_map.items()}
            display_text = id_to_display_text.get(doc_id_to_select)
            if display_text:
                new_index = self.doc_selector.findText(display_text)
        if new_index == -1 and self.doc_selector.count() > 0:
            new_index = 0
        self.doc_selector.setCurrentIndex(new_index)
        self.handle_document_switch(new_index)

    def _import_and_add_document(self, participant_id, title, content):
        try:
            new_doc_id = database.add_document(
                self.project_id, title, content, participant_id
            )
            self.is_dirty = False
            self.save_button.setEnabled(False)
            self.document_added.emit(new_doc_id)
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to import file: {e}")

    def delete_current_document(self):
        if not self.current_document_id:
            QMessageBox.warning(self, "No document", "No document is selected.")
            return
        selected_text = self.doc_selector.currentText()
        reply = QMessageBox.question(
            self,
            "Confirm Delete",
            f"Are you sure you want to permanently delete '{selected_text}' and all its coded segments?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No,
        )
        if reply == QMessageBox.StandardButton.Yes:
            self.is_dirty = False
            database.delete_document(self.current_document_id)
            self.document_deleted.emit()

    def on_text_changed(self):
        if not self.text_edit.isReadOnly():
            self.is_dirty = True
            self.save_button.setEnabled(True)

    def save_document(self, show_success_prompt=True):
        if not self.is_dirty or not self.current_document_id:
            return
        database.update_document_text_only(
            self.current_document_id, self.text_edit.toPlainText()
        )
        self.is_dirty = False
        self.save_button.setEnabled(False)
        if show_success_prompt:
            QMessageBox.information(
                self, "Success", "Document text saved successfully."
            )

    def handle_document_switch(self, new_index):
        if self.is_dirty and self.current_document_id is not None:
            msg_box = QMessageBox(self)
            msg_box.setWindowTitle("Unsaved Changes")
            msg_box.setText("You have unsaved changes. What would you like to do?")
            msg_box.setStandardButtons(
                QMessageBox.StandardButton.Save
                | QMessageBox.StandardButton.Discard
                | QMessageBox.StandardButton.Cancel
            )
            msg_box.setDefaultButton(QMessageBox.StandardButton.Save)
            reply = msg_box.exec()
            if reply == QMessageBox.StandardButton.Save:
                self.save_document(show_success_prompt=False)
            elif reply == QMessageBox.StandardButton.Cancel:
                id_to_display_text = {v: k for k, v in self.documents_map.items()}
                old_display_text = id_to_display_text.get(self.current_document_id)
                if old_display_text:
                    old_index = self.doc_selector.findText(old_display_text)
                    self.doc_selector.blockSignals(True)
                    self.doc_selector.setCurrentIndex(old_index)
                    self.doc_selector.blockSignals(False)
                return
        self.load_document_content()

    def load_document_content(self):
        QApplication.setOverrideCursor(Qt.CursorShape.WaitCursor)
        try:
            try:
                self.text_edit.textChanged.disconnect(self.on_text_changed)
            except RuntimeError:
                pass
            self.text_edit.setDocument(QTextDocument(self))
            self.is_dirty = False
            self.save_button.setEnabled(False)
            selected_display_text = self.doc_selector.currentText()
            if not selected_display_text:
                self.current_document_id = None
                self.current_participant_id = None
                self._coded_segments_cache = []
                self.text_edit.setReadOnly(True)
                self.text_edit.clear()
                self.import_button.setStyleSheet(
                    "QPushButton { border: 2px solid #0078d7; }"
                )
                self.word_count_label.setText("Word Count: 0")
                self.segment_count_label.setText("Coded Segments: 0")
            else:
                self.text_edit.setReadOnly(False)
                self.import_button.setStyleSheet("")
                self.text_edit.setAlignment(Qt.AlignmentFlag.AlignLeft)
                self.current_document_id = self.documents_map[selected_display_text]
                content, participant_id = database.get_document_content(
                    self.current_document_id
                )
                self.current_participant_id = participant_id
                self.text_edit.setPlainText(content)
                self.apply_all_highlights()
                word_count = len(content.split())
                self.word_count_label.setText(f"Word Count: {word_count}")

            if self._pending_highlight:
                start, end = self._pending_highlight
                self._select_and_scroll(start, end)
                self._pending_highlight = None

            self.text_edit.textChanged.connect(self.on_text_changed)
        finally:
            QApplication.restoreOverrideCursor()

    def apply_all_highlights(self):
        self.text_edit.blockSignals(True)
        original_position = self.text_edit.textCursor().position()
        try:
            cursor = self.text_edit.textCursor()
            cursor.select(QTextCursor.SelectionType.Document)
            cursor.setCharFormat(QTextCharFormat())
            cursor.clearSelection()
            self.text_edit.setTextCursor(cursor)
            if not self.current_document_id:
                self.segment_count_label.setText("Coded Segments: 0")
                return
            self._coded_segments_cache = database.get_coded_segments_for_document(
                self.current_document_id
            )
            self.segment_count_label.setText(
                f"Coded Segments: {len(self._coded_segments_cache)}"
            )
            for segment in self._coded_segments_cache:
                self.highlight_text(
                    segment["segment_start"],
                    segment["segment_end"],
                    segment["node_color"],
                )
        finally:
            cursor = self.text_edit.textCursor()
            cursor.setPosition(original_position)
            self.text_edit.setTextCursor(cursor)
            self.text_edit.blockSignals(False)

    def highlight_text(self, start, end, color_hex):
        cursor = self.text_edit.textCursor()
        cursor.setPosition(start)
        cursor.setPosition(end, QTextCursor.MoveMode.KeepAnchor)
        fmt = QTextCharFormat()
        bg_color = QColor(color_hex)
        brightness = (
            bg_color.red() * 299 + bg_color.green() * 587 + bg_color.blue() * 114
        ) / 1000
        text_color = QColor("black") if brightness > 128 else QColor("white")
        fmt.setBackground(bg_color)
        fmt.setForeground(text_color)
        cursor.setCharFormat(fmt)
        cursor.clearSelection()
        self.text_edit.setTextCursor(cursor)

    def on_cursor_position_changed(self):
        pos = self.text_edit.textCursor().position()
        found_segment = None
        for segment in self._coded_segments_cache:
            if segment["segment_start"] <= pos < segment["segment_end"]:
                found_segment = segment
                break
        if found_segment:
            self.segment_clicked.emit(found_segment["id"])
            self.node_clicked_in_content.emit(found_segment["node_id"])
            if found_segment.get("participant_id"):
                self.participant_highlight_requested.emit(
                    found_segment["participant_id"]
                )

    def on_segment_coded(self):
        self.segments_changed.emit()


class AssignParticipantDialog(QDialog):
    def __init__(self, participants, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Assign Participant")
        self.participants = {p["name"]: p["id"] for p in participants}
        layout = QVBoxLayout(self)
        label = QLabel("Assign this document to which participant?")
        self.combo = QComboBox()
        self.combo.addItems(sorted(self.participants.keys()))
        button_box = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel
        )
        button_box.accepted.connect(self.accept)
        button_box.rejected.connect(self.reject)
        layout.addWidget(label)
        layout.addWidget(self.combo)
        layout.addWidget(button_box)

    def get_selected_participant_id(self):
        return self.participants.get(self.combo.currentText())
