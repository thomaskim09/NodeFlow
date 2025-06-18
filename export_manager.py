from PySide6.QtWidgets import QFileDialog, QMessageBox
from docx import Document
from openpyxl.styles import Font
import json
import re
import database
import openpyxl


def export_to_word(project_id, parent_widget=None):
    """Exports the coded segments of a project to a .docx file."""
    file_path, _ = QFileDialog.getSaveFileName(
        parent_widget, "Save Word Report", "", "Word Documents (*.docx)"
    )
    if not file_path:
        return

    # --- Fetch and structure the data ---
    nodes = database.get_nodes_for_project(project_id)
    coded_segments = database.get_coded_segments_for_project(project_id)

    # Group segments by their node ID
    segments_by_node = {}
    for seg in coded_segments:
        node_id = seg["node_id"]
        if node_id not in segments_by_node:
            segments_by_node[node_id] = []
        segments_by_node[node_id].append(seg)

    # --- Build the Word Document ---
    doc = Document()
    doc.add_heading("Qualitative Analysis Report", 0)

    def write_nodes_recursively(parent_id=None, level=0, prefix=""):
        children = sorted(
            [n for n in nodes if n["parent_id"] == parent_id],
            key=lambda x: x["position"],
        )

        for i, node in enumerate(children):
            current_prefix = f"{prefix}{i + 1}."
            doc.add_heading(f"{current_prefix} {node['name']}", level=level + 1)

            if node["id"] in segments_by_node:
                for segment_data in segments_by_node[node["id"]]:
                    participant = segment_data["participant_name"] or "N/A"
                    text = segment_data["content_preview"]
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run(f"{participant}: ").bold = True
                    p.add_run(text)

            write_nodes_recursively(
                parent_id=node["id"], level=level + 1, prefix=current_prefix
            )

    write_nodes_recursively()

    # --- Save the document with error handling ---
    try:
        doc.save(file_path)
        QMessageBox.information(
            parent_widget,
            "Export Successful",
            f"Report successfully saved to:\n{file_path}",
        )
    except PermissionError:
        QMessageBox.critical(
            parent_widget,
            "Permission Denied",
            f"Could not save the file to:\n{file_path}\n\nPlease make sure you have permissions to write to this location and that the file is not currently open in another program.",
        )
    except Exception as e:
        QMessageBox.critical(
            parent_widget,
            "Export Error",
            f"An unexpected error occurred while saving the file:\n{e}",
        )


def export_to_json(project_id, parent_widget=None):
    """Exports the coded segments of a project to a .json file."""
    file_path, _ = QFileDialog.getSaveFileName(
        parent_widget, "Save JSON Export", "", "JSON Files (*.json)"
    )
    if not file_path:
        return

    nodes = database.get_nodes_for_project(project_id)
    coded_segments = database.get_coded_segments_for_project(project_id)

    segments_by_node = {}
    for seg in coded_segments:
        node_id = seg["node_id"]
        if node_id not in segments_by_node:
            segments_by_node[node_id] = []
        # Store a dictionary with participant and text
        segments_by_node[node_id].append(
            {
                "participant": seg["participant_name"] or "N/A",
                "text": seg["content_preview"],
                "document": seg["document_title"],
            }
        )

    def build_json_recursively(parent_id=None):
        children_data = []
        children = sorted(
            [n for n in nodes if n["parent_id"] == parent_id],
            key=lambda x: x["position"],
        )

        for node in children:
            node_obj = {
                "id": node["id"],
                "name": node["name"],
                "segments": segments_by_node.get(node["id"], []),
                "children": build_json_recursively(parent_id=node["id"]),
            }
            children_data.append(node_obj)
        return children_data

    json_output = build_json_recursively()

    # --- Save the JSON with error handling ---
    try:
        with open(file_path, "w", encoding="utf-8") as f:
            json.dump(json_output, f, ensure_ascii=False, indent=4)
        QMessageBox.information(
            parent_widget,
            "Export Successful",
            f"JSON data successfully saved to:\n{file_path}",
        )
    except PermissionError:
        QMessageBox.critical(
            parent_widget,
            "Permission Denied",
            f"Could not save the file to:\n{file_path}\n\nPlease make sure you have permissions to write to this location and that the file is not currently open in another program.",
        )
    except Exception as e:
        QMessageBox.critical(
            parent_widget,
            "Export Error",
            f"An unexpected error occurred while saving the file:\n{e}",
        )


def export_to_excel(project_id, parent_widget=None):
    """Exports coded segments to an Excel file with one sheet per node."""

    file_path, _ = QFileDialog.getSaveFileName(
        parent_widget, "Save Excel Report", "", "Excel Files (*.xlsx)"
    )
    if not file_path:
        return

    # Fetch data for the ENTIRE project
    nodes = database.get_nodes_for_project(project_id)
    coded_segments = database.get_coded_segments_for_project(project_id)

    # --- Data Structuring for Traversal ---
    nodes_by_id = {n["id"]: n for n in nodes}
    nodes_by_parent = {n_id: [] for n_id in nodes_by_id}
    nodes_by_parent[None] = []  # For root nodes

    for n_id, node in nodes_by_id.items():
        nodes_by_parent.get(node["parent_id"], []).append(node)

    # Sort all children lists by position
    for children_list in nodes_by_parent.values():
        children_list.sort(key=lambda x: x["position"])

    # --- Descendant and Segment Mapping ---
    descendant_cache = {}

    def get_all_descendant_ids(node_id):
        if node_id in descendant_cache:
            return descendant_cache[node_id]
        descendants = []
        for child_node in nodes_by_parent.get(node_id, []):
            descendants.append(child_node["id"])
            descendants.extend(get_all_descendant_ids(child_node["id"]))
        descendant_cache[node_id] = descendants
        return descendants

    # --- Workbook Creation ---
    wb = openpyxl.Workbook()
    if "Sheet" in wb.sheetnames:
        wb.remove(wb["Sheet"])

    def sanitize_sheet_name(name):
        name = re.sub(r"[\\/*?:\[\]]", "", name)
        return name[:31]

    # --- Recursive Sheet Creation ---
    def create_sheets_recursively(parent_id, prefix=""):
        children = nodes_by_parent.get(parent_id, [])
        for i, node in enumerate(children):
            node_id = node["id"]
            current_prefix = f"{prefix}{i + 1}."
            sheet_name = sanitize_sheet_name(f"{current_prefix} {node['name']}")

            # Create worksheet
            ws = wb.create_sheet(title=sheet_name)
            headers = ["Participant", "Coded Segment", "Document"]
            ws.append(headers)

            # --- NEW: Bold headers ---
            header_font = Font(bold=True)
            for cell in ws[1]:
                cell.font = header_font

            # Get all node IDs for this sheet (self + descendants)
            ids_to_include = [node_id] + get_all_descendant_ids(node_id)

            # Filter segments and write to sheet
            segments_for_sheet = [
                s for s in coded_segments if s["node_id"] in ids_to_include
            ]
            for seg in segments_for_sheet:
                participant = seg["participant_name"] or "N/A"
                ws.append([participant, seg["content_preview"], seg["document_title"]])

            # Adjust column widths
            ws.column_dimensions["A"].width = 25
            ws.column_dimensions["B"].width = 80
            ws.column_dimensions["C"].width = 40

            # Recurse for children
            create_sheets_recursively(node_id, prefix=current_prefix)

    # Start the process from the root (nodes with no parent)
    create_sheets_recursively(None)

    # --- Save the workbook with error handling ---
    try:
        wb.save(file_path)
        QMessageBox.information(
            parent_widget,
            "Export Successful",
            f"Excel report successfully saved to:\n{file_path}",
        )
    except PermissionError:
        QMessageBox.critical(
            parent_widget,
            "Permission Denied",
            f"Could not save the file to:\n{file_path}\n\nPlease make sure you have permissions to write to this location and that the file is not currently open in another program.",
        )
    except Exception as e:
        QMessageBox.critical(
            parent_widget,
            "Export Error",
            f"An unexpected error occurred while saving the file:\n{e}",
        )


def export_node_family_to_word(project_id, start_node_id, parent_widget=None):
    """Exports a specific node and its children to a .docx file."""
    if not start_node_id:
        return

    # Fetch all project data
    all_nodes = database.get_nodes_for_project(project_id)
    coded_segments = database.get_coded_segments_for_project(project_id)
    nodes_map = {n["id"]: n for n in all_nodes}

    start_node = nodes_map.get(start_node_id)
    if not start_node:
        return

    file_path, _ = QFileDialog.getSaveFileName(
        parent_widget,
        f"Save Word Report for '{start_node['name']}'",
        "",
        "Word Documents (*.docx)",
    )
    if not file_path:
        return

    # Filter segments for all nodes in the family first
    segments_by_node = {}
    ids_to_include = [start_node_id] + get_all_descendant_ids(
        start_node_id, nodes_map, all_nodes
    )
    for seg in coded_segments:
        if seg["node_id"] in ids_to_include:
            node_id = seg["node_id"]
            if node_id not in segments_by_node:
                segments_by_node[node_id] = []
            segments_by_node[node_id].append(seg)

    # Build the document
    doc = Document()
    doc.add_heading(f"Report for Node: {start_node['name']}", 0)

    # --- REFACTORED: Correct recursive function ---
    def write_nodes_recursively(node_id, level, prefix):
        node = nodes_map.get(node_id)
        if not node:
            return

        # 1. Write the current node's info
        doc.add_heading(f"{prefix} {node['name']}", level=level)
        if node_id in segments_by_node:
            for seg_data in segments_by_node[node_id]:
                participant = seg_data["participant_name"] or "N/A"
                text = seg_data["content_preview"]
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"{participant}: ").bold = True
                p.add_run(text)

        # 2. Recurse for children
        children = sorted(
            [n for n in all_nodes if n["parent_id"] == node_id],
            key=lambda x: x["position"],
        )
        for i, child_node in enumerate(children):
            child_prefix = f"{prefix}{i + 1}."
            write_nodes_recursively(child_node["id"], level + 1, child_prefix)

    # Start the recursion with the selected node
    write_nodes_recursively(start_node_id, level=1, prefix="1.")

    # --- Save the document with error handling ---
    try:
        doc.save(file_path)
        QMessageBox.information(
            parent_widget,
            "Export Successful",
            f"Report successfully saved to:\n{file_path}",
        )
    except PermissionError:
        QMessageBox.critical(
            parent_widget,
            "Permission Denied",
            f"Could not save the file to:\n{file_path}\n\nPlease make sure you have permissions to write to this location and that the file is not currently open in another program.",
        )
    except Exception as e:
        QMessageBox.critical(
            parent_widget,
            "Export Error",
            f"An unexpected error occurred while saving the file:\n{e}",
        )


def get_all_descendant_ids(start_node_id, nodes_map, all_nodes):
    """Helper to get all descendant IDs for a given node."""
    nodes_by_parent = {n_id: [] for n_id in nodes_map}
    nodes_by_parent[None] = []
    for n in all_nodes:
        nodes_by_parent.setdefault(n["parent_id"], []).append(n)

    descendants = []
    nodes_to_visit = [start_node_id]
    while nodes_to_visit:
        current_id = nodes_to_visit.pop(0)
        children = nodes_by_parent.get(current_id, [])
        for child in children:
            descendants.append(child["id"])
            nodes_to_visit.append(child["id"])
    return descendants


# --- NEW: Selective node family export to Excel ---
def export_node_family_to_excel(project_id, start_node_id, parent_widget=None):
    """Exports a specific node and its children to an .xlsx file."""
    if not start_node_id:
        return

    all_nodes = database.get_nodes_for_project(project_id)
    coded_segments = database.get_coded_segments_for_project(project_id)
    nodes_map = {n["id"]: n for n in all_nodes}

    start_node = nodes_map.get(start_node_id)
    if not start_node:
        return

    file_path, _ = QFileDialog.getSaveFileName(
        parent_widget,
        f"Save Excel Report for '{start_node['name']}'",
        "",
        "Excel Files (*.xlsx)",
    )
    if not file_path:
        return

    ids_to_include = [start_node_id] + get_all_descendant_ids(
        start_node_id, nodes_map, all_nodes
    )

    wb = openpyxl.Workbook()
    if "Sheet" in wb.sheetnames:
        wb.remove(wb["Sheet"])

    def sanitize_sheet_name(name):
        return re.sub(r"[\\/*?:\[\]]", "", name)[:31]

    # Create one sheet for the whole family
    sheet_name = sanitize_sheet_name(start_node["name"])
    ws = wb.create_sheet(title=sheet_name)
    headers = ["Node", "Participant", "Coded Segment", "Document"]
    ws.append(headers)

    header_font = Font(bold=True)
    for cell in ws[1]:
        cell.font = header_font

    # Filter segments and write to sheet
    segments_for_sheet = [s for s in coded_segments if s["node_id"] in ids_to_include]
    for seg in sorted(segments_for_sheet, key=lambda s: s["node_name"]):
        ws.append(
            [
                seg["node_name"],
                seg["participant_name"] or "N/A",
                seg["content_preview"],
                seg["document_title"],
            ]
        )

    ws.column_dimensions["A"].width = 30
    ws.column_dimensions["B"].width = 25
    ws.column_dimensions["C"].width = 80
    ws.column_dimensions["D"].width = 40

    # --- Save the workbook with error handling ---
    try:
        wb.save(file_path)
        QMessageBox.information(
            parent_widget,
            "Export Successful",
            f"Excel report successfully saved to:\n{file_path}",
        )
    except PermissionError:
        QMessageBox.critical(
            parent_widget,
            "Permission Denied",
            f"Could not save the file to:\n{file_path}\n\nPlease make sure you have permissions to write to this location and that the file is not currently open in another program.",
        )
    except Exception as e:
        QMessageBox.critical(
            parent_widget,
            "Export Error",
            f"An unexpected error occurred while saving the file:\n{e}",
        )


def export_node_family_to_excel_multi_sheet(
    project_id, start_node_id, parent_widget=None
):
    """Exports a specific node and its children to an .xlsx file with multiple sheets."""
    if not start_node_id:
        return

    # Fetch all project data
    all_nodes = database.get_nodes_for_project(project_id)
    coded_segments = database.get_coded_segments_for_project(project_id)
    nodes_map = {n["id"]: n for n in all_nodes}

    start_node = nodes_map.get(start_node_id)
    if not start_node:
        return

    file_path, _ = QFileDialog.getSaveFileName(
        parent_widget,
        f"Save Excel Report for '{start_node['name']}'",
        "",
        "Excel Files (*.xlsx)",
    )
    if not file_path:
        return

    # --- Data Structuring for Traversal ---
    nodes_by_parent = {n_id: [] for n_id in nodes_map}
    nodes_by_parent[None] = []
    for n in all_nodes:
        nodes_by_parent.setdefault(n["parent_id"], []).append(n)

    for children_list in nodes_by_parent.values():
        children_list.sort(key=lambda x: x["position"])

    # --- Workbook Creation ---
    wb = openpyxl.Workbook()
    if "Sheet" in wb.sheetnames:
        wb.remove(wb["Sheet"])

    def sanitize_sheet_name(name):
        return re.sub(r"[\\/*?:\[\]]", "", name)[:31]

    # --- Recursive Sheet Creation ---
    def create_sheets_recursively(node, prefix):
        node_id = node["id"]
        sheet_name = sanitize_sheet_name(f"{prefix} {node['name']}")

        ws = wb.create_sheet(title=sheet_name)
        headers = ["Participant", "Coded Segment", "Document"]
        ws.append(headers)

        header_font = Font(bold=True)
        for cell in ws[1]:
            cell.font = header_font

        ids_to_include_for_this_sheet = [node_id] + get_all_descendant_ids(
            node_id, nodes_map, all_nodes
        )

        segments_for_sheet = [
            s for s in coded_segments if s["node_id"] in ids_to_include_for_this_sheet
        ]
        for seg in segments_for_sheet:
            participant = seg["participant_name"] or "N/A"
            ws.append([participant, seg["content_preview"], seg["document_title"]])

        ws.column_dimensions["A"].width = 25
        ws.column_dimensions["B"].width = 80
        ws.column_dimensions["C"].width = 40

        # Recurse for children
        children = nodes_by_parent.get(node_id, [])
        for i, child_node in enumerate(children):
            create_sheets_recursively(child_node, f"{prefix}{i + 1}.")

    # Start the process from the start_node with prefix "1."
    create_sheets_recursively(start_node, "1.")

    # --- Save the workbook with error handling ---
    try:
        wb.save(file_path)
        QMessageBox.information(
            parent_widget,
            "Export Successful",
            f"Excel report successfully saved to:\n{file_path}",
        )
    except PermissionError:
        QMessageBox.critical(
            parent_widget,
            "Permission Denied",
            f"Could not save the file to:\n{file_path}\n\nPlease make sure you have permissions to write to this location and that the file is not currently open in another program.",
        )
    except Exception as e:
        QMessageBox.critical(
            parent_widget,
            "Export Error",
            f"An unexpected error occurred while saving the file:\n{e}",
        )
