from PySide6.QtWidgets import (
    QWidget,
    QVBoxLayout,
    QHBoxLayout,
    QComboBox,
    QPushButton,
    QTextEdit,
    QFileDialog,
    QMessageBox,
    QDialog,
    QLabel,
    QDialogButtonBox,
    QApplication,
    QFrame,
    QInputDialog,
)
from PySide6.QtCore import Qt, Signal
from PySide6.QtGui import (
    QTextCursor,
    QColor,
    QTextCharFormat,
    QTextDocument,
)
import os
import database
import docx


class ContentView(QWidget):
    document_deleted = Signal()
    segment_clicked = Signal(int)
    text_selection_changed = Signal(bool)

    def __init__(self, project_id):
        super().__init__()
        self.project_id = project_id
        self.documents_map = {}
        self.current_document_id = None
        self.current_participant_id = None
        self.is_dirty = False
        self._coded_segments_cache = []

        main_layout = QVBoxLayout(self)
        top_bar_layout = QHBoxLayout()

        title_label = QLabel("Document View")
        font = title_label.font()
        font.setBold(True)
        title_label.setFont(font)

        self.doc_selector = QComboBox()
        self.doc_selector.setMinimumWidth(300)

        self.import_button = QPushButton("⬆️")
        self.import_button.setToolTip("Import Document (.txt, .docx)")
        self.import_button.setFixedSize(28, 28)

        self.save_button = QPushButton("💾")
        self.save_button.setToolTip("Save Changes")
        self.save_button.setFixedSize(28, 28)
        self.save_button.setEnabled(False)

        delete_button = QPushButton("🗑️")
        delete_button.setToolTip("Delete Current Document")
        delete_button.setFixedSize(28, 28)

        self.text_edit = QTextEdit()
        self.text_edit.setReadOnly(False)

        info_bar = QFrame()
        info_bar.setFrameShape(QFrame.Shape.StyledPanel)
        info_bar_layout = QHBoxLayout(info_bar)
        info_bar_layout.setContentsMargins(5, 2, 5, 2)
        self.word_count_label = QLabel("Word Count: 0")
        self.segment_count_label = QLabel("Coded Segments: 0")
        info_bar_layout.addWidget(self.word_count_label)
        info_bar_layout.addStretch()
        info_bar_layout.addWidget(self.segment_count_label)

        top_bar_layout.addWidget(title_label)
        top_bar_layout.addWidget(self.doc_selector)
        top_bar_layout.addStretch()
        top_bar_layout.addWidget(self.import_button)
        top_bar_layout.addWidget(self.save_button)
        top_bar_layout.addWidget(delete_button)

        main_layout.addLayout(top_bar_layout)
        main_layout.addWidget(self.text_edit)
        main_layout.addWidget(info_bar)

        # Connect signals
        self.import_button.clicked.connect(self.import_document)
        self.save_button.clicked.connect(self.save_document)
        delete_button.clicked.connect(self.delete_current_document)
        self.doc_selector.currentIndexChanged.connect(self.handle_document_switch)
        self.text_edit.textChanged.connect(self.on_text_changed)
        self.text_edit.cursorPositionChanged.connect(self.on_cursor_position_changed)
        self.text_edit.selectionChanged.connect(self.on_selection_changed_for_coding)

        self.load_document_list()

    def on_selection_changed_for_coding(self):
        has_selection = self.text_edit.textCursor().hasSelection()
        self.text_selection_changed.emit(has_selection)

    def load_document_list(self):
        current_doc_id = self.current_document_id
        self.doc_selector.blockSignals(True)
        self.doc_selector.clear()

        docs = database.get_documents_for_project(self.project_id)
        self.documents_map = {
            f"{doc['title']} ({doc['participant_name'] or 'Unassigned'})": doc["id"]
            for doc in docs
        }
        id_to_display_text = {v: k for k, v in self.documents_map.items()}
        for display_text in sorted(self.documents_map.keys()):
            self.doc_selector.addItem(display_text)

        self.doc_selector.blockSignals(False)
        new_text_to_select = id_to_display_text.get(current_doc_id)
        index = self.doc_selector.findText(new_text_to_select)

        if index != -1:
            self.doc_selector.setCurrentIndex(index)
        elif self.doc_selector.count() > 0:
            self.doc_selector.setCurrentIndex(0)
        else:
            self.doc_selector.setCurrentIndex(-1)
            self.load_document_content()

    def _import_and_add_document(self, participant_id, file_path):
        """Helper function to read a file and add it to the database."""
        try:
            content = ""
            if file_path.lower().endswith(".docx"):
                doc = docx.Document(file_path)
                content = "\n\n".join([p.text for p in doc.paragraphs])
            else:
                with open(file_path, "r", encoding="utf-8") as f:
                    content = f.read()

            title = os.path.basename(file_path)
            new_doc_id = database.add_document(
                self.project_id, title, content, participant_id
            )
            self.load_document_list()

            new_doc_display_text = next(
                (k for k, v in self.documents_map.items() if v == new_doc_id), None
            )
            if new_doc_display_text:
                self.doc_selector.setCurrentText(new_doc_display_text)

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to import file: {e}")

    def import_document(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Import Document", "", "Documents (*.txt *.docx)"
        )
        if not file_path:
            return

        participants = database.get_participants_for_project(self.project_id)
        if not participants:
            name, ok = QInputDialog.getText(
                self,
                "Create First Participant",
                "No participants exist. Enter a name for the first participant:",
            )
            if ok and name.strip():
                database.add_participant(self.project_id, name.strip())
                participants = database.get_participants_for_project(self.project_id)
                parent_workspace = self.parent().parent().parent()
                parent_workspace.participant_manager.load_participants()
            else:
                return

        if len(participants) == 1:
            participant_id = participants[0]["id"]
            self._import_and_add_document(participant_id, file_path)
        else:
            dialog = AssignParticipantDialog(participants, self)
            if dialog.exec() == QDialog.DialogCode.Accepted:
                participant_id = dialog.get_selected_participant_id()
                if participant_id:
                    self._import_and_add_document(participant_id, file_path)

    def delete_current_document(self):
        if not self.current_document_id:
            QMessageBox.warning(self, "No document", "No document is selected.")
            return

        selected_text = self.doc_selector.currentText()
        reply = QMessageBox.question(
            self,
            "Confirm Delete",
            f"Are you sure you want to permanently delete '{selected_text}' and all its coded segments?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No,
        )

        if reply == QMessageBox.StandardButton.Yes:
            doc_id_to_delete = self.current_document_id
            self.is_dirty = False
            database.delete_document(doc_id_to_delete)
            self.load_document_list()
            self.document_deleted.emit()

    def on_text_changed(self):
        if not self.text_edit.isReadOnly():
            self.is_dirty = True
            self.save_button.setEnabled(True)

    def save_document(self):
        """Saves only the text content of the document."""
        if not self.is_dirty or not self.current_document_id:
            return

        new_content = self.text_edit.toPlainText()
        database.update_document_text_only(self.current_document_id, new_content)

        self.is_dirty = False
        self.save_button.setEnabled(False)

        QMessageBox.information(self, "Success", "Document text saved successfully.")

    def handle_document_switch(self, new_index):
        """Checks for unsaved changes before loading a new document."""
        if self.is_dirty and self.current_document_id is not None:
            msg_box = QMessageBox(self)
            msg_box.setWindowTitle("Unsaved Changes")
            msg_box.setText("You have unsaved changes. What would you like to do?")
            msg_box.setStandardButtons(
                QMessageBox.StandardButton.Save
                | QMessageBox.StandardButton.Discard
                | QMessageBox.StandardButton.Cancel
            )
            msg_box.setDefaultButton(QMessageBox.StandardButton.Save)
            reply = msg_box.exec()

            if reply == QMessageBox.StandardButton.Save:
                self.save_document()
            elif reply == QMessageBox.StandardButton.Cancel:
                id_to_display_text = {v: k for k, v in self.documents_map.items()}
                old_display_text = id_to_display_text.get(self.current_document_id)
                if old_display_text:
                    old_index = self.doc_selector.findText(old_display_text)
                    self.doc_selector.blockSignals(True)
                    self.doc_selector.setCurrentIndex(old_index)
                    self.doc_selector.blockSignals(False)
                return

        self.load_document_content()

    def load_document_content(self):
        QApplication.setOverrideCursor(Qt.CursorShape.WaitCursor)
        try:
            try:
                self.text_edit.textChanged.disconnect(self.on_text_changed)
            except RuntimeError:
                pass

            self.text_edit.setDocument(QTextDocument(self))
            self.is_dirty = False
            self.save_button.setEnabled(False)

            selected_display_text = self.doc_selector.currentText()
            if not selected_display_text:
                self.current_document_id = None
                self.current_participant_id = None
                self._coded_segments_cache = []
                self.text_edit.setReadOnly(True)
                self.text_edit.clear()
                self.import_button.setStyleSheet(
                    "QPushButton { border: 2px solid #0078d7; }"
                )
                self.word_count_label.setText("Word Count: 0")
                self.segment_count_label.setText("Coded Segments: 0")
            else:
                self.text_edit.setReadOnly(False)
                self.import_button.setStyleSheet("")
                self.text_edit.setAlignment(Qt.AlignmentFlag.AlignLeft)
                self.current_document_id = self.documents_map[selected_display_text]
                content, participant_id = database.get_document_content(
                    self.current_document_id
                )
                self.current_participant_id = participant_id
                self.text_edit.setPlainText(content)
                self.apply_all_highlights()

                word_count = len(content.split())
                self.word_count_label.setText(f"Word Count: {word_count}")

            self.text_edit.textChanged.connect(self.on_text_changed)
        finally:
            QApplication.restoreOverrideCursor()

    def apply_all_highlights(self):
        """Applies highlights for all coded segments in the current document."""
        # Block signals to prevent textChanged from firing during these changes
        self.text_edit.blockSignals(True)
        try:
            if not self.current_document_id:
                return

            self._coded_segments_cache = database.get_coded_segments_for_document(
                self.current_document_id
            )
            self.segment_count_label.setText(
                f"Coded Segments: {len(self._coded_segments_cache)}"
            )

            # Clear all existing formatting before applying new highlights
            cursor = self.text_edit.textCursor()
            cursor.select(QTextCursor.SelectionType.Document)
            cursor.setCharFormat(QTextCharFormat())
            cursor.clearSelection()
            self.text_edit.setTextCursor(cursor)

            # Apply highlights for each segment
            for segment in self._coded_segments_cache:
                self.highlight_text(
                    segment["segment_start"],
                    segment["segment_end"],
                    segment["node_color"],
                )
        finally:
            # Always unblock signals, even if an error occurred
            self.text_edit.blockSignals(False)

    def highlight_text(self, start, end, color_hex):
        """Helper method to apply color formatting to a text range."""
        cursor = self.text_edit.textCursor()
        cursor.setPosition(start)
        cursor.setPosition(end, QTextCursor.MoveMode.KeepAnchor)

        fmt = QTextCharFormat()
        bg_color = QColor(color_hex)
        brightness = (
            bg_color.red() * 299 + bg_color.green() * 587 + bg_color.blue() * 114
        ) / 1000
        text_color = QColor("black") if brightness > 128 else QColor("white")

        fmt.setBackground(bg_color)
        fmt.setForeground(text_color)
        cursor.setCharFormat(fmt)

        cursor.clearSelection()
        self.text_edit.setTextCursor(cursor)

    def on_cursor_position_changed(self):
        pos = self.text_edit.textCursor().position()
        found_segment = None
        for segment in self._coded_segments_cache:
            if segment["segment_start"] <= pos < segment["segment_end"]:
                found_segment = segment
                break

        if found_segment:
            self.segment_clicked.emit(found_segment["id"])


class AssignParticipantDialog(QDialog):
    def __init__(self, participants, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Assign Participant")
        self.participants = {p["name"]: p["id"] for p in participants}
        layout = QVBoxLayout(self)
        label = QLabel("Assign this document to which participant?")
        self.combo = QComboBox()
        self.combo.addItems(sorted(self.participants.keys()))
        button_box = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel
        )
        button_box.accepted.connect(self.accept)
        button_box.rejected.connect(self.reject)
        layout.addWidget(label)
        layout.addWidget(self.combo)
        layout.addWidget(button_box)

    def get_selected_participant_id(self):
        name = self.combo.currentText()
        return self.participants.get(name)
